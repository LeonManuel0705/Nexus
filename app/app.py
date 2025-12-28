from flask import Flask, request, jsonify, send_from_directory, send_file, render_template
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import os
import threading
from datetime import datetime
from io import BytesIO
import json
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.colors import HexColor

import database as db
from recorder import get_recorder, list_audio_devices, get_default_input_device
from transcriber import (
    cleanup_audio_file, get_available_models,
    preload_model, transcribe_audio_chunk, get_ollama_status,
    get_correction_status, reset_session, get_text_corrector,
    apply_ai_correction
)
from voice_profile import get_profile_manager, verify_teacher_audio
from quality_feedback import (
    get_feedback_collector, analyze_transcription, save_feedback,
    apply_logic_correction, get_analyzer
)
from adaptive_system import (
    get_adaptive_system, set_teacher as adaptive_set_teacher,
    record_feedback as adaptive_record_feedback, get_param,
    get_current_params, get_teacher_stats, get_change_log
)

app = Flask(__name__, static_folder='static')
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

realtime_active = False
realtime_thread = None
correction_thread = None
current_transcription = []
corrected_segments = []
last_correction_index = 0
correction_lock = threading.Lock()
current_language = "de"

voice_filter_enabled = False
voice_filter_profile_id = None
voice_filter_mode = "include"
enrollment_active = False
chunk_errors = []

CORRECTION_DELAY_SECONDS = 10

models_ready = {
    'whisper': False,
    'speaker': False,
    'mlx': False,
    'languagetool': False
}
models_progress = {
    'whisper': {'status': 'pending', 'progress': 0, 'message': 'Waiting...'},
    'speaker': {'status': 'pending', 'progress': 0, 'message': 'Waiting...'},
    'mlx': {'status': 'pending', 'progress': 0, 'message': 'Waiting...'},
    'languagetool': {'status': 'pending', 'progress': 0, 'message': 'Waiting...'}
}
init_complete = False
init_lock = threading.Lock()

def initialize_all_models():
    global models_ready, models_progress, init_complete

    def update_progress(model, status, progress, message):
        with init_lock:
            models_progress[model] = {'status': status, 'progress': progress, 'message': message}
            socketio.emit('model_progress', {
                'model': model,
                'status': status,
                'progress': progress,
                'message': message,
                'all_progress': models_progress
            })

    update_progress('whisper', 'loading', 10, 'Loading Whisper model...')
    try:
        preload_model()
        models_ready['whisper'] = True
        update_progress('whisper', 'ready', 100, 'Whisper ready')
    except Exception as e:
        update_progress('whisper', 'error', 0, f'Error: {str(e)[:50]}')

    update_progress('speaker', 'loading', 10, 'Loading speaker recognition...')
    try:
        manager = get_profile_manager()
        manager._load_model()
        models_ready['speaker'] = True
        update_progress('speaker', 'ready', 100, 'Speaker model ready')
    except Exception as e:
        update_progress('speaker', 'error', 0, f'Error: {str(e)[:50]}')

    update_progress('mlx', 'loading', 10, 'Loading AI corrector (MLX)...')
    try:
        from ai_corrector import get_corrector
        corrector = get_corrector()
        corrector._load_model()
        if corrector._mlx_available:
            models_ready['mlx'] = True
            update_progress('mlx', 'ready', 100, 'MLX AI corrector ready')
        else:
            models_ready['mlx'] = True
            update_progress('mlx', 'skipped', 100, 'MLX not available (using LanguageTool)')
    except Exception as e:
        models_ready['mlx'] = True
        update_progress('mlx', 'skipped', 100, f'MLX skipped: {str(e)[:30]}')

    update_progress('languagetool', 'loading', 10, 'Loading LanguageTool...')
    try:
        from ai_corrector import get_corrector
        corrector = get_corrector()
        corrector._get_language_tool('de')
        models_ready['languagetool'] = True
        update_progress('languagetool', 'ready', 100, 'LanguageTool ready')
    except Exception as e:
        models_ready['languagetool'] = True
        update_progress('languagetool', 'skipped', 100, f'LanguageTool skipped')

    with init_lock:
        init_complete = True

    socketio.emit('models_ready', {
        'ready': True,
        'models': models_ready,
        'progress': models_progress
    })
    print("All models initialized!")

threading.Thread(target=initialize_all_models, daemon=True).start()

def analyze_audio_quality(audio_data, sample_rate):
    import numpy as np

    if len(audio_data) == 0:
        return {'score': 0, 'issues': ['no_audio'], 'energy': 0, 'snr': 0}

    energy = float(np.sqrt(np.mean(audio_data ** 2)))

    if energy < 0.001:
        return {'score': 10, 'issues': ['too_quiet'], 'energy': energy, 'snr': 0}

    issues = []
    score = 100

    if energy < 0.01:
        issues.append('low_volume')
        score -= 20
    elif energy > 0.8:
        issues.append('clipping')
        score -= 30

    zero_crossings = np.sum(np.abs(np.diff(np.signbit(audio_data))))
    zcr = zero_crossings / len(audio_data)

    if zcr > 0.3:
        issues.append('high_noise')
        score -= 25

    fft = np.fft.rfft(audio_data)
    freqs = np.fft.rfftfreq(len(audio_data), 1/sample_rate)
    magnitude = np.abs(fft)

    speech_band = (freqs >= 300) & (freqs <= 3400)
    noise_band = (freqs < 300) | (freqs > 3400)

    speech_energy = np.sum(magnitude[speech_band] ** 2) if np.any(speech_band) else 0
    noise_energy = np.sum(magnitude[noise_band] ** 2) if np.any(noise_band) else 1

    snr = 10 * np.log10(speech_energy / noise_energy + 1e-10) if noise_energy > 0 else 0

    if snr < 5:
        issues.append('poor_snr')
        score -= 20
    elif snr < 10:
        issues.append('moderate_snr')
        score -= 10

    score = max(0, min(100, score))

    return {
        'score': int(score),
        'issues': issues,
        'energy': round(float(energy), 4),
        'snr': round(float(snr), 2),
        'zcr': round(float(zcr), 4)
    }

def create_word_document(note):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading(note['title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(41, 98, 255)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_run = meta_para.add_run(f"Created: {note['created_at']}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_run.font.italic = True

    if note.get('language'):
        lang_display = "German" if note['language'] == 'de' else "English"
        meta_para.add_run("  •  ")
        lang_run = meta_para.add_run(f"Language: {lang_display}")
        lang_run.font.size = Pt(10)
        lang_run.font.color.rgb = RGBColor(128, 128, 128)
        lang_run.font.italic = True

    if note.get('audio_duration') and note['audio_duration'] > 0:
        mins = int(note['audio_duration'] // 60)
        secs = int(note['audio_duration'] % 60)
        meta_para.add_run("  •  ")
        dur_run = meta_para.add_run(f"Duration: {mins:02d}:{secs:02d}")
        dur_run.font.size = Pt(10)
        dur_run.font.color.rgb = RGBColor(128, 128, 128)
        dur_run.font.italic = True

    if note.get('tags') and len(note['tags']) > 0:
        doc.add_paragraph()
        tags_para = doc.add_paragraph()
        tags_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tags_label = tags_para.add_run("Tags: ")
        tags_label.font.bold = True
        tags_label.font.size = Pt(10)
        tags_label.font.color.rgb = RGBColor(100, 100, 100)

        for i, tag in enumerate(note['tags']):
            tag_run = tags_para.add_run(f" {tag} ")
            tag_run.font.size = Pt(9)
            tag_run.font.color.rgb = RGBColor(255, 255, 255)
            tag_run.font.bold = True
            if i < len(note['tags']) - 1:
                tags_para.add_run("  ")

    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    doc.add_paragraph()

    content_heading = doc.add_heading("Transcription", level=1)
    for run in content_heading.runs:
        run.font.color.rgb = RGBColor(60, 60, 60)

    paragraphs = note['content'].split('\n\n') if note.get('content') else ['']
    for para_text in paragraphs:
        if para_text.strip():
            content_para = doc.add_paragraph()
            content_para.paragraph_format.space_after = Pt(12)
            content_para.paragraph_format.line_spacing = 1.5

            content_run = content_para.add_run(para_text.strip())
            content_run.font.size = Pt(11)
            content_run.font.color.rgb = RGBColor(40, 40, 40)

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("─" * 30)
    footer_run.font.color.rgb = RGBColor(200, 200, 200)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2.add_run("Generated by VoiceNotes")
    footer2_run.font.size = Pt(9)
    footer2_run.font.color.rgb = RGBColor(150, 150, 150)
    footer2_run.font.italic = True

    return doc

def create_folder_word_document(folder, notes):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading(folder['name'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(41, 98, 255)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"{len(notes)} Voice Notes")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(128, 128, 128)
    sub_run.font.italic = True

    doc.add_paragraph()

    toc_heading = doc.add_heading("Contents", level=1)
    for run in toc_heading.runs:
        run.font.color.rgb = RGBColor(60, 60, 60)

    for i, note in enumerate(notes, 1):
        toc_item = doc.add_paragraph()
        toc_run = toc_item.add_run(f"{i}. {note['title']}")
        toc_run.font.size = Pt(10)
        toc_run.font.color.rgb = RGBColor(80, 80, 80)
        date_run = toc_item.add_run(f"  ({note['created_at']})")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    for i, note in enumerate(notes):
        note_title = doc.add_heading(note['title'], level=1)
        for run in note_title.runs:
            run.font.color.rgb = RGBColor(41, 98, 255)

        meta_para = doc.add_paragraph()
        date_run = meta_para.add_run(f"Created: {note['created_at']}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        date_run.font.italic = True

        if note.get('audio_duration') and note['audio_duration'] > 0:
            mins = int(note['audio_duration'] // 60)
            secs = int(note['audio_duration'] % 60)
            meta_para.add_run("  •  ")
            dur_run = meta_para.add_run(f"Duration: {mins:02d}:{secs:02d}")
            dur_run.font.size = Pt(10)
            dur_run.font.color.rgb = RGBColor(128, 128, 128)
            dur_run.font.italic = True

        if note.get('tags') and len(note['tags']) > 0:
            tags_para = doc.add_paragraph()
            tags_label = tags_para.add_run("Tags: ")
            tags_label.font.bold = True
            tags_label.font.size = Pt(10)
            for tag in note['tags']:
                tag_run = tags_para.add_run(f" {tag} ")
                tag_run.font.size = Pt(9)
                tag_run.font.color.rgb = RGBColor(41, 98, 255)

        doc.add_paragraph()

        paragraphs = note['content'].split('\n\n') if note.get('content') else ['']
        for para_text in paragraphs:
            if para_text.strip():
                content_para = doc.add_paragraph()
                content_para.paragraph_format.space_after = Pt(12)
                content_para.paragraph_format.line_spacing = 1.5
                content_run = content_para.add_run(para_text.strip())
                content_run.font.size = Pt(11)
                content_run.font.color.rgb = RGBColor(40, 40, 40)

        if i < len(notes) - 1:
            doc.add_paragraph()
            divider = doc.add_paragraph()
            divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
            div_run = divider.add_run("─" * 40)
            div_run.font.color.rgb = RGBColor(200, 200, 200)
            doc.add_paragraph()

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generated by VoiceNotes")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.font.italic = True

    return doc

def create_pdf_document(note):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=25*mm, leftMargin=25*mm,
                           topMargin=25*mm, bottomMargin=25*mm)

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=HexColor('#2962ff'),
        spaceAfter=12,
        alignment=1  # Center
    )

    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor('#808080'),
        alignment=1,
        spaceAfter=20
    )

    content_style = ParagraphStyle(
        'Content',
        parent=styles['Normal'],
        fontSize=11,
        leading=18,
        spaceAfter=12,
        textColor=HexColor('#282828')
    )

    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#969696'),
        alignment=1
    )

    story = []

    story.append(Paragraph(note['title'], title_style))

    meta_text = f"Created: {note['created_at']}"
    if note.get('language'):
        lang_display = "German" if note['language'] == 'de' else "English"
        meta_text += f" | Language: {lang_display}"
    if note.get('audio_duration') and note['audio_duration'] > 0:
        mins = int(note['audio_duration'] // 60)
        secs = int(note['audio_duration'] % 60)
        meta_text += f" | Duration: {mins:02d}:{secs:02d}"
    story.append(Paragraph(meta_text, meta_style))

    if note.get('tags') and len(note['tags']) > 0:
        tags_text = "Tags: " + ", ".join(note['tags'])
        story.append(Paragraph(tags_text, meta_style))

    story.append(Spacer(1, 20))

    content = note.get('content', '')
    paragraphs = content.split('\n\n') if content else ['']
    for para_text in paragraphs:
        if para_text.strip():
            safe_text = para_text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_text, content_style))

    story.append(Spacer(1, 40))
    story.append(Paragraph("Generated by VoiceNotes", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer

def create_folder_pdf_document(folder, notes):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=25*mm, leftMargin=25*mm,
                           topMargin=25*mm, bottomMargin=25*mm)

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        textColor=HexColor('#2962ff'),
        spaceAfter=12,
        alignment=1
    )

    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=HexColor('#808080'),
        alignment=1,
        spaceAfter=30
    )

    note_title_style = ParagraphStyle(
        'NoteTitle',
        parent=styles['Heading2'],
        fontSize=18,
        textColor=HexColor('#2962ff'),
        spaceBefore=20,
        spaceAfter=8
    )

    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor('#808080'),
        spaceAfter=12
    )

    content_style = ParagraphStyle(
        'Content',
        parent=styles['Normal'],
        fontSize=11,
        leading=18,
        spaceAfter=12,
        textColor=HexColor('#282828')
    )

    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#969696'),
        alignment=1
    )

    story = []

    story.append(Paragraph(folder['name'], title_style))
    story.append(Paragraph(f"{len(notes)} Voice Notes", subtitle_style))

    for i, note in enumerate(notes):
        if i > 0:
            story.append(Spacer(1, 20))

        story.append(Paragraph(note['title'], note_title_style))

        meta_text = f"Created: {note['created_at']}"
        if note.get('audio_duration') and note['audio_duration'] > 0:
            mins = int(note['audio_duration'] // 60)
            secs = int(note['audio_duration'] % 60)
            meta_text += f" | Duration: {mins:02d}:{secs:02d}"
        story.append(Paragraph(meta_text, meta_style))

        content = note.get('content', '')
        paragraphs = content.split('\n\n') if content else ['']
        for para_text in paragraphs:
            if para_text.strip():
                safe_text = para_text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_text, content_style))

    story.append(Spacer(1, 40))
    story.append(Paragraph("Generated by VoiceNotes", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer

@app.route('/')
def home():
    return render_template('home.html', active_page='home')

@app.route('/app')
def application():
    return render_template('app.html')

@app.route('/app/features')
def features():
    return render_template('features.html', active_page='features')

@app.route('/app/how-it-works')
def how_it_works():
    return render_template('how_it_works.html', active_page='how-it-works')

@app.route('/app/docs')
def docs():
    return render_template('docs.html', active_page='docs')

@app.route('/app/about')
def about():
    return render_template('about.html', active_page='about')

@app.route('/app/terms')
def terms():
    return render_template('terms.html', active_page='terms')

@app.route('/static/<path:path>')
def serve_static(path):
    return send_from_directory('static', path)

def background_correction_worker():
    global realtime_active, current_transcription, corrected_segments, last_correction_index, current_language

    corrector = get_text_corrector(current_language)

    while realtime_active:
        time.sleep(2)

        if not realtime_active:
            break

        with correction_lock:
            total_segments = len(current_transcription)
            safe_index = max(0, total_segments - 3)

            if safe_index > last_correction_index:
                segments_to_correct = current_transcription[last_correction_index:safe_index]

                if segments_to_correct:
                    text_to_correct = ' '.join(segments_to_correct)
                    corrected_text = corrector.correct_full(text_to_correct)
                    corrected_text = apply_ai_correction(corrected_text, current_language)

                    corrected_segments.append(corrected_text)
                    last_correction_index = safe_index

                    full_corrected = ' '.join(corrected_segments)
                    remaining = ' '.join(current_transcription[safe_index:])
                    if remaining:
                        full_corrected += ' ' + remaining

                    socketio.emit('correction_update', {
                        'corrected_text': full_corrected,
                        'segments_corrected': safe_index,
                        'total_segments': total_segments
                    })

def realtime_transcription_worker():
    global realtime_active, current_transcription, current_language, voice_filter_mode

    recorder = get_recorder()
    corrector = get_text_corrector(current_language)

    voice_manager = get_profile_manager() if voice_filter_enabled else None
    skipped_chunks = 0
    accepted_chunks = 0
    similarity_history = []
    audio_quality_scores = []

    while realtime_active and recorder.is_recording():
        chunk = recorder.get_next_chunk(timeout=0.5)

        if chunk is not None:
            try:
                audio_quality = analyze_audio_quality(chunk, recorder.sample_rate)
                audio_quality_scores.append(audio_quality['score'])

                if voice_filter_enabled and voice_manager is not None:
                    is_match, confidence = voice_manager.verify_audio(chunk)
                    similarity = voice_manager._current_profile.verify_multi(
                        voice_manager._extract_embedding(chunk), top_k=5
                    )[1] if voice_manager._current_profile else 0

                    similarity_history.append(confidence)
                    if len(similarity_history) > 100:
                        similarity_history.pop(0)

                    dashboard_data = {
                        'similarity': float(confidence),
                        'threshold': float(get_param('voice_threshold')),
                        'is_match': bool(is_match),
                        'mode': voice_filter_mode,
                        'accepted': int(accepted_chunks),
                        'rejected': int(skipped_chunks),
                        'audio_quality': audio_quality,
                        'similarity_history': [float(x) for x in similarity_history[-20:]],
                        'avg_quality': float(sum(audio_quality_scores[-20:]) / len(audio_quality_scores[-20:])) if audio_quality_scores else 0.0
                    }
                    socketio.emit('voice_filter_dashboard', dashboard_data)

                    if voice_filter_mode == "include":
                        if not is_match:
                            skipped_chunks += 1
                            socketio.emit('voice_filter_status', {
                                'filtering': True,
                                'mode': 'include',
                                'skipped': int(skipped_chunks),
                                'accepted': int(accepted_chunks),
                                'last_confidence': float(confidence)
                            })
                            continue
                        accepted_chunks += 1
                    else:
                        if is_match:
                            skipped_chunks += 1
                            socketio.emit('voice_filter_status', {
                                'filtering': True,
                                'mode': 'exclude',
                                'skipped': int(skipped_chunks),
                                'accepted': int(accepted_chunks),
                                'last_confidence': float(confidence)
                            })
                            continue
                        accepted_chunks += 1

                raw_text, corrected_text, confidence = transcribe_audio_chunk(
                    chunk,
                    sample_rate=recorder.sample_rate,
                    language=current_language,
                    apply_instant_correction=True
                )

                if corrected_text and corrected_text.strip():
                    logic_corrected, changes = apply_logic_correction(corrected_text.strip(), current_language)

                    with correction_lock:
                        current_transcription.append(logic_corrected)
                        full_text = ' '.join(current_transcription)

                    full_corrected = corrector.correct_instant(full_text)
                    full_corrected, _ = apply_logic_correction(full_corrected, current_language)

                    socketio.emit('transcription_update', {
                        'text': logic_corrected,
                        'raw_text': raw_text.strip() if raw_text else '',
                        'full_text': full_corrected,
                        'confidence': confidence,
                        'is_final': False,
                        'is_corrected': True,
                        'voice_verified': voice_filter_enabled,
                        'logic_corrections': len(changes)
                    })

            except Exception as e:
                print(f"Transcription error: {e}")
                import traceback
                traceback.print_exc()

@socketio.on('start_realtime')
def handle_start_realtime(data):
    global realtime_active, realtime_thread, correction_thread
    global current_transcription, corrected_segments, last_correction_index, current_language

    folder_id = data.get('folder_id')
    current_language = data.get('language', 'de')
    recorder = get_recorder()

    if recorder.is_recording():
        emit('error', {'message': 'Already recording'})
        return

    with correction_lock:
        current_transcription = []
        corrected_segments = []
        last_correction_index = 0

    realtime_active = True
    reset_session()

    subject = None
    if folder_id:
        folder = db.get_folder(folder_id)
        if folder:
            subject = folder.get('name')

    recorder.start_recording()

    realtime_thread = threading.Thread(target=realtime_transcription_worker, daemon=True)
    realtime_thread.start()

    correction_thread = threading.Thread(target=background_correction_worker, daemon=True)
    correction_thread.start()

    emit('recording_started', {
        'message': 'Recording started',
        'language': current_language,
        'subject': subject
    })

@socketio.on('stop_realtime')
def handle_stop_realtime(data):
    global realtime_active, current_transcription, corrected_segments, last_correction_index, current_language

    folder_id = data.get('folder_id')
    recorder = get_recorder()

    if not recorder.is_recording():
        emit('error', {'message': 'Not recording'})
        return

    realtime_active = False
    audio_path, duration = recorder.stop_recording()

    emit('recording_stopping', {'message': 'Processing and saving...'})

    raw_segments = list(current_transcription)
    corrected_segs = list(corrected_segments)
    last_idx = last_correction_index
    lang = current_language

    def finalize_and_save():
        with correction_lock:
            if last_idx < len(raw_segments):
                remaining_segments = raw_segments[last_idx:]
                if remaining_segments:
                    corrector = get_text_corrector(lang)
                    remaining_text = ' '.join(remaining_segments)
                    corrected_remaining = corrector.correct_full(remaining_text)
                    corrected_remaining = apply_ai_correction(corrected_remaining, lang)
                    corrected_segs.append(corrected_remaining)

            final_text = ' '.join(corrected_segs)

        if not final_text.strip():
            final_text = ' '.join(raw_segments)
            if final_text.strip():
                corrector = get_text_corrector(lang)
                final_text = corrector.correct_full(final_text)
                final_text = apply_ai_correction(final_text, lang)

        if audio_path:
            cleanup_audio_file(audio_path)

        if final_text.strip():
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            title = f"Voice Note - {timestamp}"

            note_id = db.create_note(
                title=title,
                content=final_text,
                folder_id=folder_id,
                language=lang,
                audio_duration=duration
            )

            note = db.get_note(note_id)

            socketio.emit('recording_stopped', {
                'success': True,
                'note': note,
                'final_text': final_text,
                'duration': duration
            })
        else:
            socketio.emit('recording_stopped', {
                'success': False,
                'message': 'No speech detected'
            })

    threading.Thread(target=finalize_and_save, daemon=True).start()

@socketio.on('set_language')
def handle_set_language(data):
    global current_language
    current_language = data.get('language', 'de')
    emit('language_changed', {'language': current_language})

@app.route('/api/recording/start', methods=['POST'])
def start_recording():
    recorder = get_recorder()
    if recorder.is_recording():
        return jsonify({'success': False, 'error': 'Already recording'}), 400

    reset_session()
    success = recorder.start_recording()
    return jsonify({'success': success})

@app.route('/api/recording/stop', methods=['POST'])
def stop_recording():
    global current_language, current_transcription
    recorder = get_recorder()
    if not recorder.is_recording():
        return jsonify({'success': False, 'error': 'Not recording'}), 400

    data = request.get_json() or {}
    folder_id = data.get('folder_id')
    language = data.get('language', current_language)

    audio_path, duration = recorder.stop_recording()

    if not audio_path:
        return jsonify({'success': False, 'error': 'No audio recorded'}), 400

    corrector = get_text_corrector(language)
    final_text = corrector.correct_full(' '.join(current_transcription))
    cleanup_audio_file(audio_path)

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    title = f"Voice Note - {timestamp}"

    note_id = db.create_note(
        title=title,
        content=final_text,
        folder_id=folder_id,
        language=language,
        audio_duration=duration
    )

    note = db.get_note(note_id)

    return jsonify({
        'success': True,
        'note': note,
        'transcription': {
            'text': final_text,
            'language': language,
            'duration': duration
        }
    })

@app.route('/api/recording/status', methods=['GET'])
def recording_status():
    recorder = get_recorder()
    return jsonify({
        'recording': recorder.is_recording(),
        'duration': recorder.get_duration()
    })

@app.route('/api/audio/devices', methods=['GET'])
def audio_devices():
    devices = list_audio_devices()
    default = get_default_input_device()
    return jsonify({'devices': devices, 'default': default})

@app.route('/api/folders', methods=['GET'])
def get_folders():
    parent_id = request.args.get('parent_id', type=int)
    if request.args.get('all') == 'true':
        folders = db.get_all_folders()
    else:
        folders = db.get_folders(parent_id)
    return jsonify({'folders': folders})

@app.route('/api/folders', methods=['POST'])
def create_folder():
    data = request.get_json()
    if not data or not data.get('name'):
        return jsonify({'error': 'Name is required'}), 400

    folder_id = db.create_folder(name=data['name'], parent_id=data.get('parent_id'))
    folder = db.get_folder(folder_id)
    return jsonify({'folder': folder}), 201

@app.route('/api/folders/<int:folder_id>', methods=['GET'])
def get_folder(folder_id):
    folder = db.get_folder(folder_id)
    if not folder:
        return jsonify({'error': 'Folder not found'}), 404
    return jsonify({'folder': folder})

@app.route('/api/folders/<int:folder_id>', methods=['PUT'])
def update_folder(folder_id):
    data = request.get_json()
    if not data or not data.get('name'):
        return jsonify({'error': 'Name is required'}), 400

    success = db.rename_folder(folder_id, data['name'])
    if not success:
        return jsonify({'error': 'Folder not found'}), 404

    folder = db.get_folder(folder_id)
    return jsonify({'folder': folder})

@app.route('/api/folders/<int:folder_id>', methods=['DELETE'])
def delete_folder(folder_id):
    success = db.delete_folder(folder_id)
    if not success:
        return jsonify({'error': 'Folder not found'}), 404
    return jsonify({'success': True})

@app.route('/api/folders/<int:folder_id>/subfolders', methods=['GET'])
def get_subfolders(folder_id):
    subfolders = db.get_folders(folder_id)
    return jsonify({'folders': subfolders})

@app.route('/api/notes', methods=['GET'])
def get_notes():
    folder_id = request.args.get('folder_id', type=int)
    notes = db.get_notes(folder_id)
    return jsonify({'notes': notes})

@app.route('/api/notes/<int:note_id>', methods=['GET'])
def get_note(note_id):
    note = db.get_note(note_id)
    if not note:
        return jsonify({'error': 'Note not found'}), 404
    return jsonify({'note': note})

@app.route('/api/notes/<int:note_id>', methods=['PUT'])
def update_note(note_id):
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    kwargs = {
        'title': data.get('title'),
        'content': data.get('content')
    }
    if 'folder_id' in data:
        kwargs['folder_id'] = data['folder_id']

    success = db.update_note(note_id, **kwargs)

    if not success:
        return jsonify({'error': 'Note not found or no changes'}), 404

    note = db.get_note(note_id)
    return jsonify({'note': note})

@app.route('/api/notes/<int:note_id>', methods=['DELETE'])
def delete_note(note_id):
    success = db.delete_note(note_id)
    if not success:
        return jsonify({'error': 'Note not found'}), 404
    return jsonify({'success': True})

@app.route('/api/notes/search', methods=['GET'])
def search_notes():
    query = request.args.get('q', '')
    if not query:
        return jsonify({'notes': []})
    notes = db.search_notes(query)
    return jsonify({'notes': notes})

@app.route('/api/tags', methods=['GET'])
def get_tags():
    tags = db.get_all_tags()
    return jsonify({'tags': tags})

@app.route('/api/notes/<int:note_id>/tags', methods=['POST'])
def add_tag(note_id):
    data = request.get_json()
    if not data or not data.get('tag'):
        return jsonify({'error': 'Tag is required'}), 400

    success = db.add_tag_to_note(note_id, data['tag'])
    note = db.get_note(note_id)
    return jsonify({'note': note, 'added': success})

@app.route('/api/notes/<int:note_id>/tags/<tag_name>', methods=['DELETE'])
def remove_tag(note_id, tag_name):
    success = db.remove_tag_from_note(note_id, tag_name)
    note = db.get_note(note_id)
    return jsonify({'note': note, 'removed': success})

@app.route('/api/tags/<tag_name>/notes', methods=['GET'])
def get_notes_by_tag(tag_name):
    notes = db.get_notes_by_tag(tag_name)
    return jsonify({'notes': notes})

@app.route('/api/notes/<int:note_id>/export', methods=['GET'])
def export_note(note_id):
    format_type = request.args.get('format', 'txt')
    note = db.get_note(note_id)

    if not note:
        return jsonify({'error': 'Note not found'}), 404

    if format_type == 'docx':
        doc = create_word_document(note)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        safe_title = "".join(c for c in note['title'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_title}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    elif format_type == 'pdf':
        buffer = create_pdf_document(note)
        safe_title = "".join(c for c in note['title'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_title}.pdf"
        return send_file(buffer, as_attachment=True, download_name=filename,
                        mimetype='application/pdf')

    elif format_type == 'md':
        content = f"# {note['title']}\n\n"
        content += f"*Created: {note['created_at']}*\n\n"
        if note['tags']:
            content += f"**Tags:** {', '.join(note['tags'])}\n\n"
        content += "---\n\n"
        content += note['content']
        filename = f"{note['title']}.md"
        mimetype = 'text/markdown'
    else:
        content = f"{note['title']}\n"
        content += f"Created: {note['created_at']}\n"
        if note['tags']:
            content += f"Tags: {', '.join(note['tags'])}\n"
        content += "\n" + "=" * 40 + "\n\n"
        content += note['content']
        filename = f"{note['title']}.txt"
        mimetype = 'text/plain'

    buffer = BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype=mimetype)

@app.route('/api/folders/<int:folder_id>/export', methods=['GET'])
def export_folder(folder_id):
    format_type = request.args.get('format', 'md')
    folder = db.get_folder(folder_id)
    notes = db.get_notes(folder_id)

    if not folder:
        return jsonify({'error': 'Folder not found'}), 404

    if format_type == 'docx':
        doc = create_folder_word_document(folder, notes)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        safe_name = "".join(c for c in folder['name'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_name}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    elif format_type == 'pdf':
        buffer = create_folder_pdf_document(folder, notes)
        safe_name = "".join(c for c in folder['name'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_name}.pdf"
        return send_file(buffer, as_attachment=True, download_name=filename,
                        mimetype='application/pdf')

    elif format_type == 'md':
        content = f"# {folder['name']}\n\n"
        for note in notes:
            content += f"## {note['title']}\n\n"
            content += f"*Created: {note['created_at']}*\n\n"
            content += note['content'] + "\n\n---\n\n"
        filename = f"{folder['name']}.md"
        mimetype = 'text/markdown'
    else:
        content = f"{folder['name']}\n{'=' * 40}\n\n"
        for note in notes:
            content += f"{note['title']}\n"
            content += f"Created: {note['created_at']}\n"
            content += "-" * 20 + "\n"
            content += note['content'] + "\n\n"
        filename = f"{folder['name']}.txt"
        mimetype = 'text/plain'

    buffer = BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype=mimetype)

@app.route('/api/export/database', methods=['GET'])
def export_database():
    folders = db.get_all_folders()
    notes = db.get_notes()
    tags = db.get_all_tags()

    export_data = {
        'version': '2.0',
        'exported_at': datetime.now().isoformat(),
        'folders': folders,
        'notes': notes,
        'tags': tags
    }

    buffer = BytesIO()
    buffer.write(json.dumps(export_data, indent=2, ensure_ascii=False).encode('utf-8'))
    buffer.seek(0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"voicenotes_backup_{timestamp}.json"

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/json')

@app.route('/api/settings/models', methods=['GET'])
def get_models():
    models = get_available_models()
    return jsonify({'models': models})

@app.route('/api/settings/ollama', methods=['GET'])
def get_ollama():
    status = get_ollama_status()
    return jsonify(status)

@app.route('/api/settings/status', methods=['GET'])
def get_system_status():
    correction = get_correction_status()

    return jsonify({
        'offline_ready': True,
        'ai_correction': correction['ai_correction'],
        'speaker_diarization': correction['speaker_diarization'],
        'whisper': {
            'realtime_model': 'small',
            'processing': 'live'
        }
    })

@app.route('/api/models/check', methods=['GET'])
def check_models():
    from model_manager import check_all_models
    return jsonify(check_all_models())

@app.route('/api/models/download/<model_name>', methods=['POST'])
def start_model_download(model_name):
    from model_manager import download_model_async

    valid_models = ['whisper', 'mlx', 'speaker']
    if model_name not in valid_models:
        return jsonify({'error': 'Invalid model name'}), 400

    download_model_async(model_name)
    return jsonify({'status': 'started', 'model': model_name})

@app.route('/api/models/progress', methods=['GET'])
def get_model_progress():
    from model_manager import get_download_progress
    return jsonify({'downloads': get_download_progress()})

@app.route('/api/voice-profiles', methods=['GET'])
def list_voice_profiles():
    manager = get_profile_manager()
    profiles = manager.list_profiles()
    current = manager.get_current_profile()
    return jsonify({
        'profiles': profiles,
        'current_profile': current,
        'filter_enabled': voice_filter_enabled
    })

@app.route('/api/voice-profiles/current', methods=['GET'])
def get_current_voice_profile():
    manager = get_profile_manager()
    current = manager.get_current_profile()
    return jsonify({
        'profile': current,
        'filter_enabled': voice_filter_enabled
    })

@app.route('/api/voice-profiles/<profile_id>', methods=['DELETE'])
def delete_voice_profile(profile_id):
    manager = get_profile_manager()
    success = manager.delete_profile(profile_id)
    if success:
        return jsonify({'success': True})
    return jsonify({'error': 'Profile not found'}), 404

@app.route('/api/voice-profiles/<profile_id>/activate', methods=['POST'])
def activate_voice_profile(profile_id):
    global voice_filter_enabled, voice_filter_profile_id

    manager = get_profile_manager()
    if manager.load_profile(profile_id):
        voice_filter_profile_id = profile_id
        voice_filter_enabled = True

        profile = manager.get_current_profile()
        if profile:
            adaptive_result = adaptive_set_teacher(profile.get('name', 'Unknown'), profile_id)
            return jsonify({
                'success': True,
                'profile': profile,
                'adaptive': adaptive_result
            })

        return jsonify({
            'success': True,
            'profile': profile
        })
    return jsonify({'error': 'Profile not found'}), 404

@app.route('/api/voice-profiles/deactivate', methods=['POST'])
def deactivate_voice_filter():
    global voice_filter_enabled, voice_filter_profile_id

    voice_filter_enabled = False
    voice_filter_profile_id = None
    return jsonify({'success': True})

@app.route('/api/voice-profiles/filter-status', methods=['GET'])
def get_voice_filter_status():
    manager = get_profile_manager()
    return jsonify({
        'enabled': voice_filter_enabled,
        'profile_id': voice_filter_profile_id,
        'mode': voice_filter_mode,
        'profile': manager.get_current_profile() if voice_filter_enabled else None
    })

@app.route('/api/voice-profiles/set-mode', methods=['POST'])
def set_voice_filter_mode():
    global voice_filter_mode
    data = request.get_json()
    mode = data.get('mode', 'include')
    if mode in ['include', 'exclude']:
        voice_filter_mode = mode
        return jsonify({'success': True, 'mode': mode})
    return jsonify({'error': 'Invalid mode. Use "include" or "exclude"'}), 400

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    data = request.get_json()

    note_id = data.get('note_id')
    transcription = data.get('transcription', '')
    rating = data.get('rating', 3)
    folder_id = data.get('folder_id')
    teacher_name = data.get('teacher_name')
    subject = data.get('subject')
    user_comments = data.get('user_comments')
    language = data.get('language', 'de')

    if not note_id:
        return jsonify({'error': 'note_id required'}), 400

    feedback_result = save_feedback(
        note_id=note_id,
        transcription=transcription,
        rating=rating,
        folder_id=folder_id,
        teacher_name=teacher_name,
        subject=subject,
        user_comments=user_comments,
        language=language
    )

    adaptive_result = adaptive_record_feedback(
        rating=rating,
        teacher_id=voice_filter_profile_id,
        teacher_name=teacher_name,
        quality_score=feedback_result.get('analysis', {}).get('quality_score'),
        issues=feedback_result.get('analysis', {}).get('issues', [])
    )

    return jsonify({
        'feedback': feedback_result,
        'adaptive': adaptive_result
    })

@app.route('/api/feedback/analyze', methods=['POST'])
def analyze_text():
    data = request.get_json()
    text = data.get('text', '')
    language = data.get('language', 'de')

    if not text:
        return jsonify({'error': 'text required'}), 400

    result = analyze_transcription(text, language)
    return jsonify(result)

@app.route('/api/feedback/stats', methods=['GET'])
def get_feedback_stats():
    collector = get_feedback_collector()
    return jsonify(collector.get_feedback_stats())

@app.route('/api/feedback/reports', methods=['GET'])
def get_error_reports():
    collector = get_feedback_collector()
    return jsonify({'reports': collector.get_pending_reports()})

@app.route('/api/feedback/report/<filename>', methods=['GET'])
def get_report_content(filename):
    from pathlib import Path
    reports_dir = Path(__file__).parent.parent / "error_reports"
    report_path = reports_dir / filename

    if not report_path.exists():
        return jsonify({'error': 'Report not found'}), 404

    with open(report_path, 'r', encoding='utf-8') as f:
        content = f.read()

    return jsonify({'filename': filename, 'content': content})

@app.route('/api/adaptive/params', methods=['GET'])
def get_adaptive_params():
    return jsonify({
        'params': get_current_params(),
        'teacher_id': voice_filter_profile_id
    })

@app.route('/api/adaptive/teacher-stats/<profile_id>', methods=['GET'])
def get_adaptive_teacher_stats(profile_id):
    manager = get_profile_manager()
    profiles = manager.list_profiles()

    teacher_name = "Unknown"
    for p in profiles:
        if p.get('profile_id') == profile_id:
            teacher_name = p.get('name', 'Unknown')
            break

    stats = get_teacher_stats(profile_id, teacher_name)
    return jsonify(stats)

@app.route('/api/adaptive/log', methods=['GET'])
def get_adaptive_log():
    limit = request.args.get('limit', 50, type=int)
    log = get_change_log(limit)
    return jsonify({'log': log})

@app.route('/api/adaptive/reset/<profile_id>', methods=['POST'])
def reset_adaptive_teacher(profile_id):
    manager = get_profile_manager()
    profiles = manager.list_profiles()

    teacher_name = "Unknown"
    for p in profiles:
        if p.get('profile_id') == profile_id:
            teacher_name = p.get('name', 'Unknown')
            break

    system = get_adaptive_system()
    result = system.reset_teacher(profile_id, teacher_name)
    return jsonify(result)

@app.route('/api/adaptive/threshold', methods=['POST'])
def set_threshold_live():
    data = request.get_json()
    new_threshold = data.get('threshold')

    if new_threshold is None or not (0.1 <= new_threshold <= 0.8):
        return jsonify({'error': 'Threshold must be between 0.1 and 0.8'}), 400

    system = get_adaptive_system()
    old_value = system.get_param('voice_threshold')
    system._current_params['voice_threshold'] = new_threshold
    system._log_change('voice_threshold', old_value, new_threshold, 'Manual adjustment via dashboard')

    return jsonify({
        'success': True,
        'old_threshold': old_value,
        'new_threshold': new_threshold
    })

@app.route('/api/chunk-error', methods=['POST'])
def mark_chunk_error():
    global chunk_errors
    data = request.get_json()

    error_entry = {
        'timestamp': datetime.now().isoformat(),
        'chunk_index': data.get('chunk_index'),
        'text': data.get('text', ''),
        'error_type': data.get('error_type', 'transcription_error'),
        'note_id': data.get('note_id'),
        'teacher_id': voice_filter_profile_id
    }

    chunk_errors.append(error_entry)
    if len(chunk_errors) > 100:
        chunk_errors = chunk_errors[-100:]

    system = get_adaptive_system()
    system.record_feedback(
        rating=1,
        teacher_id=voice_filter_profile_id,
        issues=[{'type': error_entry['error_type']}]
    )

    return jsonify({'success': True, 'error_logged': error_entry})

@app.route('/api/chunk-errors', methods=['GET'])
def get_chunk_errors():
    return jsonify({'errors': chunk_errors[-50:]})

@app.route('/api/models/init-status', methods=['GET'])
def get_init_status():
    with init_lock:
        return jsonify({
            'ready': init_complete,
            'models': models_ready,
            'progress': models_progress
        })

@app.route('/api/audio-quality', methods=['POST'])
def check_audio_quality():
    data = request.get_json()
    if 'audio' not in data:
        return jsonify({'error': 'audio data required'}), 400

    import numpy as np
    import base64

    audio_bytes = base64.b64decode(data['audio'])
    audio_array = np.frombuffer(audio_bytes, dtype=np.float32)
    sample_rate = data.get('sample_rate', 16000)

    quality = analyze_audio_quality(audio_array, sample_rate)
    return jsonify(quality)

@socketio.on('start_enrollment')
def handle_start_enrollment(data):
    global enrollment_active

    teacher_name = data.get('name', 'Teacher')
    manager = get_profile_manager()

    profile_id = manager.start_enrollment(teacher_name)
    enrollment_active = True

    emit('enrollment_started', {
        'profile_id': profile_id,
        'name': teacher_name,
        'min_duration': 30
    })

    recorder = get_recorder()
    print(f"[Enrollment] Checking recorder state: is_recording={recorder.is_recording()}")

    if recorder.is_recording():
        print("[Enrollment] ERROR: Already recording")
        emit('enrollment_error', {'error': 'Recording already in progress. Please stop the current recording first.'})
        manager.cancel_enrollment()
        enrollment_active = False
        return

    print("[Enrollment] Starting recording...")
    recording_started = recorder.start_recording()
    print(f"[Enrollment] Recording started: {recording_started}")

    if not recording_started:
        print(f"[Enrollment] ERROR: Failed to start - {recorder.last_error}")
        emit('enrollment_error', {'error': f'Failed to start recording: {recorder.last_error or "No audio input device found"}'})
        manager.cancel_enrollment()
        enrollment_active = False
        return

    print(f"[Enrollment] Recording active: {recorder.is_recording()}")

    def enrollment_worker():
        print("Enrollment worker started")
        chunks_processed = 0
        while enrollment_active and recorder.is_recording():
            chunk = recorder.get_enrollment_chunk(timeout=0.5)
            if chunk is not None:
                chunks_processed += 1
                print(f"Enrollment chunk {chunks_processed}: {len(chunk)} samples, {len(chunk)/16000:.1f}s")
                result = manager.add_enrollment_audio(chunk)
                print(f"Enrollment progress: {result}")
                socketio.emit('enrollment_progress', result)
        print(f"Enrollment worker stopped after {chunks_processed} chunks")

    threading.Thread(target=enrollment_worker, daemon=True).start()
    print("Enrollment worker thread started")

@socketio.on('stop_enrollment')
def handle_stop_enrollment(data):
    global enrollment_active, voice_filter_enabled, voice_filter_profile_id

    enrollment_active = False
    recorder = get_recorder()

    if recorder.is_recording():
        recorder.stop_recording()

    manager = get_profile_manager()
    result = manager.finish_enrollment()

    if result.get('status') == 'success':
        voice_filter_profile_id = result['profile_id']
        manager.load_profile(result['profile_id'])
        voice_filter_enabled = True
        result['filter_activated'] = True

    emit('enrollment_complete', result)

@socketio.on('cancel_enrollment')
def handle_cancel_enrollment():
    global enrollment_active

    enrollment_active = False
    recorder = get_recorder()

    if recorder.is_recording():
        recorder.stop_recording()

    manager = get_profile_manager()
    manager.cancel_enrollment()

    emit('enrollment_cancelled', {'success': True})

if __name__ == '__main__':
    print("\n" + "=" * 50)
    print("  Voice Notes - Real-time Transcription")
    print("  Open http://localhost:5050 in your browser")
    print("=" * 50 + "\n")
    socketio.run(app, host='127.0.0.1', port=5050, debug=False, allow_unsafe_werkzeug=True)
